"""
文件解析服务
支持DOCX、TXT、XLSX、PDF格式的简历解析
"""
import io
from typing import Optional
from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader


class FileParser:
    """文件解析器"""

    @staticmethod
    def parse_file(file_content: bytes, filename: str) -> str:
        """
        根据文件扩展名选择解析方法

        Args:
            file_content: 文件二进制内容
            filename: 文件名（用于判断扩展名）

        Returns:
            解析后的纯文本内容
        """
        ext = filename.lower().split('.')[-1] if '.' in filename else ''

        if ext == 'docx':
            return FileParser.parse_docx(file_content)
        elif ext == 'txt':
            return FileParser.parse_txt(file_content)
        elif ext == 'xlsx':
            return FileParser.parse_xlsx(file_content)
        elif ext == 'pdf':
            return FileParser.parse_pdf(file_content)
        else:
            raise ValueError(f"不支持的文件格式: {ext}")

    @staticmethod
    def parse_docx(file_content: bytes) -> str:
        """
        解析DOCX文件

        Args:
            file_content: DOCX文件二进制内容

        Returns:
            解析后的纯文本内容
        """
        try:
            doc = Document(io.BytesIO(file_content))
            paragraphs = []

            # 提取段落文本
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        paragraphs.append(' | '.join(row_text))

            return '\n'.join(paragraphs)

        except Exception as e:
            raise ValueError(f"DOCX文件解析失败: {str(e)}")

    @staticmethod
    def parse_txt(file_content: bytes) -> str:
        """
        解析TXT文件

        Args:
            file_content: TXT文件二进制内容

        Returns:
            解析后的纯文本内容
        """
        try:
            # 尝试多种编码
            encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16']

            for encoding in encodings:
                try:
                    return file_content.decode(encoding)
                except UnicodeDecodeError:
                    continue

            raise ValueError("无法识别文件编码，请使用UTF-8或GBK编码")

        except Exception as e:
            raise ValueError(f"TXT文件解析失败: {str(e)}")

    @staticmethod
    def parse_xlsx(file_content: bytes) -> str:
        """
        解析XLSX文件

        Args:
            file_content: XLSX文件二进制内容

        Returns:
            解析后的纯文本内容
        """
        try:
            wb = load_workbook(io.BytesIO(file_content))
            all_text = []

            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                sheet_text = []

                for row in sheet.iter_rows(values_only=True):
                    # 过滤空值并转为字符串
                    row_values = [str(cell) if cell is not None else '' for cell in row]
                    row_text = ' | '.join([v for v in row_values if v.strip()])
                    if row_text.strip():
                        sheet_text.append(row_text)

                if sheet_text:
                    all_text.append(f"=== {sheet_name} ===")
                    all_text.extend(sheet_text)

            return '\n'.join(all_text)

        except Exception as e:
            raise ValueError(f"XLSX文件解析失败: {str(e)}")

    @staticmethod
    def parse_pdf(file_content: bytes) -> str:
        """
        解析PDF文件

        Args:
            file_content: PDF文件二进制内容

        Returns:
            解析后的纯文本内容
        """
        try:
            pdf_reader = PdfReader(io.BytesIO(file_content))
            all_text = []

            # 提取每一页的文本
            for page_num, page in enumerate(pdf_reader.pages, 1):
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    # 清理文本：移除多余的空白字符
                    cleaned_text = '\n'.join(
                        line.strip() for line in page_text.split('\n') if line.strip()
                    )
                    if cleaned_text:
                        all_text.append(cleaned_text)

            if not all_text:
                raise ValueError("PDF文件未提取到文本内容，可能是扫描件或图片PDF")

            return '\n\n'.join(all_text)

        except ValueError:
            raise
        except Exception as e:
            raise ValueError(f"PDF文件解析失败: {str(e)}")


# 创建全局实例
parser = FileParser()
